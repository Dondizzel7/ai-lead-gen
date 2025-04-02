"""
DOCX Resume Processor for Resume Screening AI Module

This module handles the extraction and preprocessing of text from DOCX resumes.
It includes document structure analysis and text normalization.
"""

import logging
import os
import io
import re
from typing import Dict, List, Any, Optional, BinaryIO
import docx
from docx.document import Document as DocxDocument

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DOCXResumeProcessor:
    """
    DOCX Resume Processor class for extracting and preprocessing text from DOCX resumes.
    """
    
    def __init__(self):
        """Initialize the DOCX Resume Processor."""
        logger.info("DOCX Resume Processor initialized")
    
    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        logger.info(f"Extracting text from DOCX file: {file_path}")
        
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return {"success": False, "message": "File not found", "text": "", "metadata": {}}
        
        try:
            with open(file_path, 'rb') as file:
                return self.extract_text_from_binary(file)
        except Exception as e:
            logger.exception(f"Error extracting text from DOCX file: {str(e)}")
            return {"success": False, "message": str(e), "text": "", "metadata": {}}
    
    def extract_text_from_binary(self, file_obj: BinaryIO) -> Dict[str, Any]:
        """
        Extract text from a binary DOCX file object.
        
        Args:
            file_obj: Binary file object
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        logger.info("Extracting text from binary DOCX file")
        
        try:
            # Create a copy of the file in memory to avoid modifying the original
            file_copy = io.BytesIO(file_obj.read())
            file_obj.seek(0)  # Reset original file pointer
            
            # Load the document
            doc = docx.Document(file_copy)
            
            # Extract text
            full_text = self._extract_text_from_docx(doc)
            
            # Extract metadata
            metadata = self._extract_metadata(doc)
            
            # Preprocess text
            processed_text = self._preprocess_text(full_text)
            
            # Extract document structure
            structure = self._extract_document_structure(doc)
            
            return {
                "success": True,
                "text": processed_text,
                "raw_text": full_text,
                "metadata": metadata,
                "structure": structure,
                "extraction_method": "python-docx"
            }
        except Exception as e:
            logger.exception(f"Error extracting text from binary DOCX file: {str(e)}")
            return {"success": False, "message": str(e), "text": "", "metadata": {}}
    
    def _extract_text_from_docx(self, doc: DocxDocument) -> str:
        """
        Extract text from a DOCX document.
        
        Args:
            doc: DOCX document object
            
        Returns:
            Extracted text
        """
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    
    def _extract_metadata(self, doc: DocxDocument) -> Dict[str, Any]:
        """
        Extract metadata from DOCX document.
        
        Args:
            doc: DOCX document object
            
        Returns:
            Dictionary of metadata
        """
        try:
            core_properties = doc.core_properties
            
            return {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "keywords": core_properties.keywords or "",
                "created": str(core_properties.created) if core_properties.created else "",
                "modified": str(core_properties.modified) if core_properties.modified else "",
                "last_modified_by": core_properties.last_modified_by or "",
                "revision": core_properties.revision or 0,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
        except Exception as e:
            logger.warning(f"Metadata extraction error: {str(e)}")
            return {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
    
    def _extract_document_structure(self, doc: DocxDocument) -> Dict[str, Any]:
        """
        Extract structural information from DOCX document.
        
        Args:
            doc: DOCX document object
            
        Returns:
            Dictionary with document structure information
        """
        structure = {
            "headings": [],
            "paragraphs": [],
            "tables": []
        }
        
        # Extract headings and paragraphs
        for i, para in enumerate(doc.paragraphs):
            para_info = {
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "length": len(para.text)
            }
            
            # Check if paragraph is a heading
            if para.style and "Heading" in para.style.name:
                structure["headings"].append(para_info)
            
            structure["paragraphs"].append(para_info)
        
        # Extract table information
        for i, table in enumerate(doc.tables):
            table_info = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cell_count": len(table.rows) * len(table.columns)
            }
            structure["tables"].append(table_info)
        
        return structure
    
    def _preprocess_text(self, text: str) -> str:
        """
        Preprocess extracted text to improve quality.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Preprocessed text
        """
        if not text:
            return ""
        
        # Replace multiple newlines with a single newline
        processed = re.sub(r'\n{3,}', '\n\n', text)
        
        # Replace multiple spaces with a single space
        processed = re.sub(r' {2,}', ' ', processed)
        
        # Remove non-printable characters
        processed = re.sub(r'[^\x20-\x7E\n]', '', processed)
        
        return processed.strip()
    
    def identify_sections(self, text: str) -> Dict[str, str]:
        """
        Identify common resume sections in the text.
        
        Args:
            text: Preprocessed resume text
            
        Returns:
            Dictionary mapping section names to section content
        """
        # Common section headers in resumes
        section_patterns = {
            "contact_info": r"(?i)^(.*?)(CONTACT|PERSONAL|INFO|PROFILE)",
            "summary": r"(?i)^(.*?)(SUMMARY|OBJECTIVE|PROFESSIONAL SUMMARY|CAREER OBJECTIVE)",
            "experience": r"(?i)^(.*?)(EXPERIENCE|WORK|EMPLOYMENT|CAREER|PROFESSIONAL EXPERIENCE)",
            "education": r"(?i)^(.*?)(EDUCATION|ACADEMIC|QUALIFICATION|DEGREE)",
            "skills": r"(?i)^(.*?)(SKILLS|TECHNICAL|TECHNOLOGIES|COMPETENCIES|EXPERTISE)",
            "projects": r"(?i)^(.*?)(PROJECTS|PROJECT EXPERIENCE|PORTFOLIO)",
            "certifications": r"(?i)^(.*?)(CERTIFICATIONS|CERTIFICATES|ACCREDITATIONS)",
            "languages": r"(?i)^(.*?)(LANGUAGES|LANGUAGE PROFICIENCY)",
            "interests": r"(?i)^(.*?)(INTERESTS|HOBBIES|ACTIVITIES)",
            "references": r"(?i)^(.*?)(REFERENCES|REFEREES)"
        }
        
        # Split text into lines
        lines = text.split('\n')
        
        # Identify potential section headers
        section_headers = []
        for i, line in enumerate(lines):
            line = line.strip()
            if line and len(line) < 50 and line.upper() == line:
                section_headers.append((i, line))
        
        # If no section headers found, try alternative approach
        if not section_headers:
            for i, line in enumerate(lines):
                line = line.strip()
                if line and len(line) < 50 and any(re.search(pattern, line, re.IGNORECASE) for pattern in section_patterns.values()):
                    section_headers.append((i, line))
        
        # Extract sections
        sections = {}
        for i in range(len(section_headers)):
            header_idx, header = section_headers[i]
            
            # Determine section name
            section_name = "unknown"
            for name, pattern in section_patterns.items():
                if re.search(pattern, header, re.IGNORECASE):
                    section_name = name
                    break
            
            # Determine section end
            if i < len(section_headers) - 1:
                end_idx = section_headers[i + 1][0]
            else:
                end_idx = len(lines)
            
            # Extract section content
            section_content = '\n'.join(lines[header_idx + 1:end_idx]).strip()
            sections[section_name] = section_content
        
        return sections
    
    def analyze_resume_structure(self, text: str, doc_structure: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Analyze the structure of a resume.
        
        Args:
            text: Preprocessed resume text
            doc_structure: Document structure information (optional)
            
        Returns:
            Dictionary with resume structure analysis
        """
        # Identify sections
        sections = self.identify_sections(text)
        
        # Calculate basic metrics
        word_count = len(text.split())
        line_count = len(text.split('\n'))
        char_count = len(text)
        
        # Analyze section distribution
        section_distribution = {
            section: len(content.split()) / word_count if word_count > 0 else 0
            for section, content in sections.items()
        }
        
        # Detect potential issues
        issues = []
        if word_count < 200:
            issues.append("Resume is very short")
        if "experience" not in sections or not sections.get("experience"):
            issues.append("No experience section detected")
        if "education" not in sections or not sections.get("education"):
            issues.append("No education section detected")
        if "skills" not in sections or not sections.get("skills"):
            issues.append("No skills section detected")
        
        # Analyze formatting if document structure is provided
        formatting_analysis = {}
        if doc_structure:
            # Analyze heading usage
            heading_count = len(doc_structure.get("headings", []))
            formatting_analysis["heading_count"] = heading_count
            
            # Analyze paragraph length distribution
            para_lengths = [p.get("length", 0) for p in doc_structure.get("paragraphs", [])]
            if para_lengths:
                formatting_analysis["avg_paragraph_length"] = sum(para_lengths) / len(para_lengths)
                formatting_analysis["max_paragraph_length"] = max(para_lengths)
                formatting_analysis["min_paragraph_length"] = min(para_lengths)
            
            # Analyze table usage
            table_count = len(doc_structure.get("tables", []))
            formatting_analysis["table_count"] = table_count
            
            # Detect formatting issues
            if heading_count == 0:
                issues.append("No headings used in document")
            if table_count > 5:
                issues.append("Excessive use of tables")
        
        return {
            "sections": sections,
            "metrics": {
                "word_count": word_count,
                "line_count": line_count,
                "char_count": char_count,
                "section_count": len(sections)
            },
            "section_distribution": section_distribution,
            "formatting": formatting_analysis,
            "issues": issues
        }

# Example usage
if __name__ == "__main__":
    processor = DOCXResumeProcessor()
    
    # Example: Process a DOCX resume
    result = processor.extract_text_from_file("example_resume.docx")
    
    if result["success"]:
        print(f"Extraction method: {result['extraction_method']}")
        print(f"Paragraph count: {result['metadata'].get('paragraph_count', 'Unknown')}")
        print(f"Table count: {result['metadata'].get('table_count', 'Unknown')}")
        print(f"Word count: {len(result['text'].split())}")
        
        # Analyze resume structure
        structure = processor.analyze_resume_structure(result["text"], result.get("structure"))
        
        print("\nDetected Sections:")
        for section, content in structure["sections"].items():
            print(f"- {section.upper()}: {len(content.split())} words")
        
        print("\nFormatting Analysis:")
        for key, value in structure.get("formatting", {}).items():
            print(f"- {key}: {value}")
        
        print("\nPotential Issues:")
        for issue in structure["issues"]:
            print(f"- {issue}")
    else:
        print(f"Error: {result['message']}")
